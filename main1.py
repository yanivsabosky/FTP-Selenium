from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
import utilitis
from datetime import datetime
from docx import Document
import time
import socket
import os
import subprocess
import threading
import argparse

# This section sets up command-line argument parsing:
# 1. It defines a dictionary of valid URLs for easy lookup.
# 2. It uses argparse for robust and user-friendly command-line interface.
# 3. It validates the input and provides help information if needed.
URLS = {'b':'www.bbc.com', 'ab':'abcnews.go.com', 'g':'www.theguardian.com'}
parser = argparse.ArgumentParser(
                prog = 'News To Document',
                description = 'this program combine web scraping, document creation and file transfer functionality',
                epilog = 'hope U like our work')

parser.add_argument('-site', '--s', required=True, choices=URLS.keys(),
                    help='choose site to scrape: "b" for bbc, "ab" for abc and "g" for theguardian')
args = parser.parse_args()

#choosing from one of the option if its not in him
if args.s not in URLS.keys():
    print(parser.format_help())
    exit()

# This code starts a server in a separate process and thread:
# 1. It uses subprocess to run the server script independently.
# 2. It uses threading to not block the main program execution.
# 3. The sleep ensures the server has time to start before proceeding.
def run_server():
    global server_process
    server_process = subprocess.Popen(['python3', 'server.py'])

# Start the server in a separate thread
server_thread = threading.Thread(target=run_server)
server_thread.start()
time.sleep(2)


# This function sends a file over a network socket:
# 1. It uses a context manager (with) for proper socket handling.
# 2. It sends metadata (filename, file size) separately before the file content.
# 3. Small delays ensure separate transmission of metadata.
# 4. It reads and sends the entire file at once, which is simple but may not be suitable for very large files.
def send_file(filename, file_path):
    host = '127.0.0.1'  # The server's hostname or IP address
    port = 8080  # The port used by the server

    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
        s.connect((host, port))

        # Send filename
        s.send(filename.encode())
        time.sleep(0.1)  # Small delay to ensure separate transmission

        # Get file size and send it
        file_size = os.path.getsize(file_path)
        s.send(str(file_size).encode())
        time.sleep(0.1)  # Small delay to ensure separate transmission

        # Send file data
        with open(file_path, 'rb') as file:
            s.sendall(file.read())

    print(f"File sent to server: {filename}")

# Configure Chrome options to run headless
# the reason i did it if the user run the process in the terminal ->
# so i dont wont him to display any ui when the process is running
options = Options()
options.add_argument("--headless")  # Run in headless mode
options.add_argument("--disable-gpu")  # Disable GPU hardware acceleration
options.add_argument("--no-sandbox")  # Bypass OS security model
options.add_argument("--disable-dev-shm-usage")  # Overcome limited resource problems

url = URLS[args.s]
if url:
    # Initialize the Chrome WebDriver with the correct class and options
    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=options)
    driver.get(f'https://{url}')
    time.sleep(5)  # Reduced wait time, adjust if needed

    # Find the first three headings
    headings = utilitis.find_headings(driver, 3)

    # Find the first three paragraphs
    paragraphs = utilitis.find_paragraphs(driver, 3)

    # Create a new Document object
    doc = Document()

    # Add the current date and time
    current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    doc.add_paragraph(f"Date and Time: {current_datetime}")

    for i in range(min(len(headings), len(paragraphs))):
        doc.add_heading(f"\nHeading {i + 1}: {headings[i]}")
        doc.add_paragraph(f"Paragraph {i + 1}: {paragraphs[i]}")

    # Get the filename from the user
    filename = 'news.docx'

    # Save the document temporarily
    temp_file_path = filename
    doc.save(temp_file_path)

    # Send the file to the server
    try:
        send_file(filename, temp_file_path)
        print(f"File '{filename}' sent to server successfully.")
        print(f"The Word file '{filename}' has been created and saved on the FTP server.")
        print(f"You can find it in your Documents/received_files directory.")
    except Exception as e:
        print(f"Error sending file to server: {e}")
    finally:
        # Clean up the temporary file
        os.remove(temp_file_path)

    # Close the WebDriver
    driver.quit()

else:
    print("No URL provided")

# Terminate the server subprocess
if 'server_process' in globals():
    server_process.terminate()
    server_process.wait()

# Wait for the server thread to finish
server_thread.join()

print("Process completed. Server and thread terminated.")